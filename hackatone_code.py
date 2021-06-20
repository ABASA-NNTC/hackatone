import pymorphy2 #для лематизации
import docx #чтоб документы как-то читались
import re #для небольшой фильтрации
import os #чтоб было удобно ходить по файлам
from nltk.corpus import stopwords #тоже для фильтрации
import matplotlib.pyplot as plt
import numpy as np
# Массивы с ключевыми словами
#main - главные ключевые слова, остальные для улучшения

# Искусственный интеллект

x1_main = ['искусственный', 'интеллект', 'машинный', 'обучение']
x1_main_l = len(x1_main)
x1 = ['алгоритм', 'бот', 'анализ', 'данные', 'самообучение', 'язык', 'программирование', 'python']
x1_l = len(x1)

# Новые производственные технологии

x2_main = ['новый', 'производственный', 'технология']
x2_main_l = len(x2_main)
x2 = ['государственный', 'прибыльный', 'производство', 'проект']
x2_l = len(x2)

# Роботехника и сенсорика

x3_main = ['роботехника', 'робот', 'сенсорика']
x3_main_l = len(x3_main)
x3 = ['техника', 'электроника', 'автоматизированный', 'радиотехника', 'проектирование']
x3_l = len(x3)

# Интернет вещей

x4_main = ['интернет', 'вещь']
x4_main_l = len(x4_main)
x4 = ['сеть', 'покупка', 'продажа', 'предприятие', 'прибыль', 'бизнес']
x4_l = len(x4)

# Мобильные сети связи пятого поколения (цифровые сервисы)

x5_main = ['мобильный', 'сеть', 'пятый', 'поколение']
x5_main_l = len(x5_main)
x5 = ['цифровой', 'связь', 'интернет', 'преимущество', 'технология']
x5_l = len(x5)

# Новые коммуникационные интернет-технологии

x6_main = ['новый', 'коммуникационный', 'интернет-технология']
x6_main_l = len(x6_main)
x6 = ['интернет', 'технология', 'коммуникация', 'связь', 'преимущество']
x6_l = len(x6)

# Технологии виртуальной и дополненной реальности

x7_main = ['технология', 'виртуальный', 'vr', 'реальность', 'дополненный', 'ar']
x7_main_l = len(x7_main)
x7 = ['очки', 'погружение', 'реалистичность', 'ощущение', 'возможность']
x7_l = len(x7)

# Технологии распределенных реестров

x8_main = ['технология', 'распределённый', 'реестр']
x8_main_l = len(x8_main)
x8 = ['электронный', 'система', 'база', 'данные', 'сетевой', 'устройство', 'интернет']
x8_l = len(x8)

# Квантовые коммуникации

x9_main = ['квантовый', 'коммуникация']
x9_main_l = len(x9_main)
x9 = ['область', 'знание', 'местоположение', 'удаленный', 'технология', 'коммуникация']
x9_l = len(x9)

# Квантовые сенсоры

x10_main = ['квантовый', 'сенсор']
x10_main_l = len(x10_main)
x10 = ['технология', 'точность', 'устройство', 'измерение', 'система', 'вычисление']
x10_l = len(x10)

# Квантовые вычисления

x11_main = ['квантовый', 'вычисление']
x11_main_l = len(x11_main)
x11 = ['технология', 'точность', 'измерение', 'квант', 'компьютер']
x11_l = len(x11)



#обработка docx файла
#возвращает множество слов используемых в тексте
def convert_docx(path):
    morph = pymorphy2.MorphAnalyzer() #объект для лемматизации
    try:
        dfile = docx.Document(path) #сам документ
        all_paras = dfile.paragraphs #текст параграфов
        sod = ''  # содержание всего документа идёт сюда
        #здесь оно закидывается
        for para in all_paras:
            sod += para.text.lower() + '\n'
        for table in dfile.tables:
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    sod += cell.text.lower() + '\n'
        sod = re.sub(r'\W', ' ', sod) #удаление лишних знаков
        words = sod.split() #разделяем по словам
        norm_words = [] #слова после лемматизации пойдут сюда
        stop = stopwords.words('russian') #список стоп слов
        for word in words:
            if not word.isnumeric() and word not in stop and len(word) != 1: #не берём числа, стоп-слова и соединения из 1 буквы
                p = morph.parse(word)[0] #здесь процесс лемматизации
                norm_words.append(p.normal_form)
    except: #исключения нужны, ибо иногда
        pass
    return list(set(norm_words))

def analyze(n, file):
    sphere = ['Искусственный интеллект', 'Новые производственные технологии', 'Роботехника и сенсорика',
              "Интернет вещей", \
              "Мобильный сети связи пятого поколения(ЦС)", "Новые коммуникационные интернет-технологии", \
              "Технологии виртуальной и дополненной реальности", "Технологии распределенных реестров", \
              "Квантовые коммуникации", "Квантовые сенсоры", "Квантовые вычисления"
              ]
    k = 0
    if n == 0:
        main = x1_main
        vtor = x1
    elif n == 1:
        main = x2_main
        vtor = x2
    elif n == 2:
        main = x3_main
        vtor = x3
    elif n == 3:
        main = x4_main
        vtor = x4
    elif n == 4:
        main = x5_main
        vtor = x5
    elif n == 5:
        main = x6_main
        vtor = x6
    elif n == 6:
        main = x7_main
        vtor = x7
    elif n == 7:
        main = x8_main
        vtor = x8
    elif n == 8:
        main = x9_main
        vtor = x9
    elif n == 9:
        main = x10_main
        vtor = x10
    elif n == 10:
        main = x11_main
        vtor = x11




    # print(obl)
    ii_o = convert_docx(file)
    # print(ii_o)
    ii_o_l = (len(ii_o))  # количество слов в описании
    for i in range(ii_o_l):  # проверка на ключевые слова
        if ii_o[i] in main:
            k = k + 5
            # print('Найдено ключевое слово:', ii_o[i])

    if k == 0:
        print('Не найдено ключевых слов')

    if k >= 5:
        for j in range(ii_o_l):  # проверка на второстепенные слова
            if ii_o[j] in vtor:
                k = k + 1
                # print('Найдено второстепенное слово:', ii_o[j])

    k_max = (1 * len(main) * 5) + (1 * len(vtor) * 1)
    print(sphere[n])
    print('Вы набрали', k, 'баллов из', k_max, 'баллов\n')
    score[sphere[n]] = round(100 / k_max * k, 2)





os.chdir(r'C:\Users\vladt\Downloads\Программы ПЦС 2020')
folder = os.getcwd()
paths = []
score = {}


for root, dirs, files in os.walk(folder):
    for file in files:
        if not file.startswith('~') and file.endswith('.docx'):
            paths.append(os.path.join(root, file))
k = 0
for path in paths:
    print(path)
    filename = path[path.rfind('\\') + 1:-5]
    for i in range(11):
        analyze(i, path)
    print(path)
    for key, value in score.items():
        print(key, '-', str(value) + '%')
    print()

    pos = np.arange(len(score.keys()))
    width = 0.9  # gives histogram aspect to the bar diagram

    ax = plt.axes()
    ax.set_xticks(pos + (width / 2))
    ax.set_xticklabels(range(1,12))
    plt.axis([0, 11, 0, 100])
    plt.bar(list(score.keys()), score.values(), width, color='g')
    k += 1
    plt.savefig(r'C:\Users\vladt\Downloads\saved_figure{0}.png'.format(filename))
    plt.show()
    score = {}

